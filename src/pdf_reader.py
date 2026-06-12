"""
CV dosyalarından metin çıkarma — PDF, Word ve düz metin destekler.
"""

from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(pdf_path: str | Path) -> str:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"Dosya bulunamadı: {path}")

    reader = PdfReader(str(path))
    pages = [page.extract_text() for page in reader.pages if page.extract_text()]
    return "\n\n".join(pages)


def _iter_block_items(parent):
    """Belgedeki paragraf ve tabloları sırayla döndürür."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in parent.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _extract_textbox_texts(doc) -> list[str]:
    """Metin kutularındaki (text box) metni XML üzerinden çeker."""
    from docx.oxml.ns import qn

    texts = []
    # w:txbxContent tüm belgede aranır
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for para in txbx.iter(qn("w:p")):
            text = "".join(
                t.text for t in para.iter(qn("w:t")) if t.text
            ).strip()
            if text:
                texts.append(text)
    return texts


def extract_text_from_docx(docx_path: str | Path) -> str:
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Dosya bulunamadı: {path}")

    doc = Document(str(path))
    parts = []

    # Paragraf ve tablolar — belgedeki sıraya göre
    for block in _iter_block_items(doc):
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            seen_cells: set = set()
            for row in block.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen_cells:
                        seen_cells.add(cell_text)
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append("  |  ".join(row_texts))

    # Metin kutuları (iletişim bilgileri burada olabilir)
    textbox_texts = _extract_textbox_texts(doc)
    if textbox_texts:
        parts.insert(0, "\n".join(textbox_texts))  # Başa ekle — genelde üstte olur

    # Header / Footer
    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

    return "\n".join(parts)


def extract_text_smart(input_data: str | Path) -> str:
    """
    Dosya uzantısına göre otomatik metin çıkarır.
    Dosya yolu değilse düz metin olarak döndürür.
    """
    input_str = str(input_data)

    if len(input_str) > 500 or "\n" in input_str:
        return input_str

    path = Path(input_str)
    if not path.exists():
        return input_str

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in {".docx", ".doc"}:
        return extract_text_from_docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")

    raise ValueError(f"Desteklenmeyen format: {path.suffix}")
